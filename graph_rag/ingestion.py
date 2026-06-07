"""Async ingestion and adaptive semantic chunking."""

from __future__ import annotations

import asyncio
import json
import logging
import os
import re
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Any

import fitz
import pandas as pd

from graph_rag.config import RAGConfig
from graph_rag.models import DocumentChunk
from graph_rag.utils import CacheManager, normalize_text

logger = logging.getLogger(__name__)


class AdaptiveSemanticChunker:
    """Create chunks at semantic boundaries instead of fixed character windows."""

    def __init__(self, config: RAGConfig):
        self.config = config

    def split(self, text: str) -> list[str]:
        """Split text by headings, pages, paragraphs and sentences with soft size targets."""
        text = normalize_text(text)
        if not text:
            return []
        raw_sections = re.split(r"(?=\bPage\s+\d+:)|\n\s*\n|(?=^\s*#{1,4}\s+)", text, flags=re.MULTILINE)
        units: list[str] = []
        for section in raw_sections:
            section = section.strip()
            if not section:
                continue
            if len(section) <= self.config.chunk_max_chars:
                units.append(section)
                continue
            units.extend(s.strip() for s in re.split(r"(?<=[.!?])\s+", section) if s.strip())

        chunks: list[str] = []
        current: list[str] = []
        current_len = 0
        for unit in units:
            unit_len = len(unit)
            should_flush = (
                current
                and current_len >= self.config.chunk_min_chars
                and current_len + unit_len > self.config.chunk_target_chars
            ) or (current_len + unit_len > self.config.chunk_max_chars)
            if should_flush:
                chunks.append(" ".join(current).strip())
                current = current[-self.config.chunk_overlap_sentences :] if self.config.chunk_overlap_sentences else []
                current_len = sum(len(x) for x in current)
            current.append(unit)
            current_len += unit_len
        if current:
            chunks.append(" ".join(current).strip())
        return [c for c in chunks if c]


class GenericFileProcessor:
    """Multi-format document text extraction with lightweight metadata."""

    def process_file(self, file_path: str | Path) -> dict[str, Any]:
        """Extract content, file type and document metadata."""
        path = Path(file_path)
        ext = path.suffix.lower()
        try:
            if ext == ".pdf":
                return self._process_pdf(path)
            if ext in {".txt", ".md", ".rtf"}:
                return self._process_text(path, ext.lstrip("."))
            if ext == ".docx":
                return self._process_docx(path)
            if ext == ".json":
                return self._process_json(path)
            if ext in {".html", ".htm", ".xml"}:
                return self._process_text(path, ext.lstrip("."))
            if ext == ".odt":
                return self._process_odt(path)
        except Exception as exc:
            logger.warning("Failed to process %s: %s", path, exc)
        return {"content": "", "file_type": ext.lstrip(".") or "unknown", "metadata": {"error": "unsupported_or_empty"}}

    def _process_pdf(self, path: Path) -> dict[str, Any]:
        doc = fitz.open(path)
        pages = [f"Page {i + 1}:\n{page.get_text()}" for i, page in enumerate(doc) if page.get_text().strip()]
        doc.close()
        text = "\n\n".join(pages)
        return {"content": text, "file_type": "pdf", "metadata": {"pages": len(pages), "total_chars": len(text)}}

    def _process_text(self, path: Path, file_type: str) -> dict[str, Any]:
        for enc in ("utf-8", "latin-1"):
            try:
                text = path.read_text(encoding=enc, errors="replace")
                return {"content": text, "file_type": file_type, "metadata": {"encoding": enc, "total_chars": len(text)}}
            except UnicodeDecodeError:
                continue
        return {"content": "", "file_type": file_type, "metadata": {"error": "encoding_not_detected"}}

    def _process_json(self, path: Path) -> dict[str, Any]:
        data = json.loads(path.read_text(encoding="utf-8"))
        text = json.dumps(data, indent=2, ensure_ascii=False)
        return {"content": text, "file_type": "json", "metadata": {"total_chars": len(text)}}

    def _process_docx(self, path: Path) -> dict[str, Any]:
        from docx import Document

        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return {"content": text, "file_type": "docx", "metadata": {"paragraphs": len(doc.paragraphs)}}

    def _process_odt(self, path: Path) -> dict[str, Any]:
        with zipfile.ZipFile(path, "r") as archive:
            root = ET.fromstring(archive.read("content.xml"))
        text = " ".join(el.text for el in root.iter() if el.text)
        return {"content": text, "file_type": "odt", "metadata": {"total_chars": len(text)}}

    def detect_document_type(self, text: str, filename: str) -> str:
        """Classify common procurement and business document types."""
        haystack = f"{filename} {text[:3000]}".lower()
        type_keywords = {
            "grn": ["goods receipt", "grn", "material receipt"],
            "invoice": ["tax invoice", "sales invoice", "invoice"],
            "purchase_order": ["purchase order", "p.o.", "po number"],
            "proforma_invoice": ["proforma", "pro forma"],
            "quotation": ["quotation", "quote"],
            "delivery_note": ["delivery note", "challan"],
            "statement": ["statement", "ledger"],
        }
        for doc_type, keywords in type_keywords.items():
            if any(keyword in haystack for keyword in keywords):
                return doc_type
        return "document"

    def extract_structured_data(self, text: str) -> dict[str, Any]:
        """Extract high-value fields for graph links and source citations."""
        patterns = {
            "document_number": [r"(?:GRN|Invoice|PO|PI|Quote|Doc)\s*(?:No|Number|Code)?[:\s#-]*([A-Z0-9][A-Z0-9/-]{3,})"],
            "date": [r"(?:Date|Dated)[:\s]*(\d{1,2}[-/]\d{1,2}[-/]\d{2,4})", r"(\d{1,2}\s+\w+\s+\d{4})"],
            "amount": [r"(?:Total|Amount|Grand Total|Net Payable)[:\s]*(?:Rs\.?|INR)?\s*(\d[\d,]*\.?\d*)"],
            "tax_amount": [r"(?:GST|Tax|VAT|IGST|CGST|SGST)[:\s]*(\d[\d,]*\.?\d*)"],
            "supplier_name": [r"(?:Supplier|Vendor|From)[:\s]*([A-Za-z0-9\s&.,()-]{3,80})"],
            "customer_name": [r"(?:Customer|Buyer|Bill To|To)[:\s]*([A-Za-z0-9\s&.,()-]{3,80})"],
            "address": [r"(?:Address|Addr)[:\s]*([^\n]{10,160})"],
        }
        extracted: dict[str, Any] = {}
        for field, field_patterns in patterns.items():
            for pattern in field_patterns:
                match = re.search(pattern, text, flags=re.IGNORECASE)
                if match:
                    extracted[field] = normalize_text(match.group(1))
                    break
        return extracted


class DocumentIngestor:
    """Load spreadsheets and documents with async processing and incremental caches."""

    def __init__(self, config: RAGConfig, cache: CacheManager):
        self.config = config
        self.cache = cache
        self.processor = GenericFileProcessor()
        self.chunker = AdaptiveSemanticChunker(config)

    def load_excel_data(self, excel_path: str | Path) -> pd.DataFrame:
        """Load all xlsx/csv files into one normalized dataframe."""
        key = f"excel_{self.cache.corpus_key(excel_path)}"
        cached = self.cache.get(key)
        if cached is not None:
            return cached
        path = Path(excel_path)
        frames: list[pd.DataFrame] = []
        files = [path] if path.is_file() else [Path(root) / f for root, _, names in os.walk(path) for f in names]
        for file_path in files:
            if file_path.suffix.lower() == ".csv":
                frame = pd.read_csv(file_path)
            elif file_path.suffix.lower() in {".xlsx", ".xls"}:
                frame = pd.read_excel(file_path)
            else:
                continue
            frame["_source_file"] = file_path.name
            frames.append(frame)
        if not frames:
            return pd.DataFrame()
        df = pd.concat(frames, ignore_index=True)
        for col in df.select_dtypes(include="object").columns:
            df[col] = df[col].apply(normalize_text)
        self.cache.set(key, df)
        return df

    async def load_documents(self, folder: str | Path) -> list[DocumentChunk]:
        """Load non-spreadsheet documents concurrently."""
        path = Path(folder)
        key = f"documents_{self.cache.corpus_key(path)}"
        cached = self.cache.get(key)
        if cached is not None:
            return [DocumentChunk.from_dict(item) for item in cached]
        files = [
            Path(root) / filename
            for root, _, filenames in os.walk(path)
            for filename in filenames
            if Path(filename).suffix.lower() not in {".xlsx", ".xls", ".csv", ".pkl"}
        ]
        tasks = [asyncio.to_thread(self._process_document, file_path) for file_path in files]
        nested = await asyncio.gather(*tasks) if tasks else []
        chunks = [chunk for group in nested for chunk in group]
        self.cache.set(key, [chunk.to_dict() for chunk in chunks])
        return chunks

    def _process_document(self, path: Path) -> list[DocumentChunk]:
        data = self.processor.process_file(path)
        text = data.get("content", "")
        if not text.strip():
            return []
        doc_type = self.processor.detect_document_type(text, path.name)
        structured = self.processor.extract_structured_data(text)
        metadata = {
            **dict(data.get("metadata") or {}),
            "source_path": str(path),
            "file_type": data.get("file_type", "unknown"),
            "document_type": doc_type,
        }
        chunks = []
        for idx, chunk_text in enumerate(self.chunker.split(text)):
            chunks.append(
                DocumentChunk(
                    id=f"{path.name}_{idx}",
                    text=chunk_text,
                    source=path.name,
                    type="document",
                    metadata=metadata,
                    document_type=doc_type,
                    chunk_index=idx,
                    structured_data=structured,
                )
            )
        return chunks

    def create_excel_chunks(self, df: pd.DataFrame) -> list[DocumentChunk]:
        """Convert spreadsheet rows into retrievable chunks."""
        chunks: list[DocumentChunk] = []
        if df.empty:
            return chunks
        for idx, row in df.iterrows():
            row_data = {k: v for k, v in row.to_dict().items() if k != "_source_file" and pd.notna(v) and str(v).strip()}
            text = " | ".join(f"{k}: {v}" for k, v in row_data.items())
            if text:
                chunks.append(
                    DocumentChunk(
                        id=f"excel_row_{idx}",
                        text=text,
                        source=str(row.get("_source_file", "excel")),
                        type="excel",
                        metadata={"source_file": row.get("_source_file", "excel"), "document_type": "spreadsheet_row"},
                        document_type="spreadsheet_row",
                        row_index=int(idx),
                        row_data=row_data,
                    )
                )
        return chunks
